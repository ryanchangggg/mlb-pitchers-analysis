"""
Generate a formal .docx report comparing Ohtani, Sanchez, and Misiorowski (2026 season).
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ---- Styles ----
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ---- Title ----
title = doc.add_heading('Comparative Analysis of Three National League Aces', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_heading('Ohtani, Sanchez, and Misiorowski — 2026 Season', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
   'This report analyzes Statcast pitch-by-pitch data for three elite NL pitchers '
   'in the 2026 season. The goal is to evaluate whether Shohei Ohtani\'s pitching '
   'performance is genuinely superior to that of Cristopher Sanchez and Jacob Misiorowski.'
)

# ---- Section 1: Data Overview ----
doc.add_heading('1. Data Overview', level=2)
doc.add_paragraph(
   'The dataset contains 4,797 total pitch events. Ohtani has thrown 1,335 pitches '
   'across 14 games, Sanchez 1,800 pitches in 19 games, and Misiorowski 1,662 pitches '
   'in 18 games. Ohtani\'s reduced workload is immediately notable; he has pitched '
   'approximately 35% fewer pitches than Sanchez.'
)

# ---- Section 2: Pitch Arsenal ----
doc.add_heading('2. Pitch Arsenal Analysis', level=2)
doc.add_paragraph(
   'Ohtani deploys the widest arsenal (7 pitch types), but variety does not equate to '
   'superiority. Sanchez uses only 3 pitch types (sinker, changeup, slider) yet achieves '
   'elite results through sheer command and deception. Misiorowski relies heavily on a '
   'dominant four-seam fastball (63.1% usage) averaging 100.5 mph.'
)

# Insert velocity comparison table
table = doc.add_table(rows=4, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Metric', 'Ohtani', 'Sanchez', 'Misiorowski']
for i, h in enumerate(headers):
   table.rows[0].cells[i].text = h
data_rows = [
   ['Primary FB Velo', '98.1 mph (FF)', '95.2 mph (SI)', '100.5 mph (FF)'],
   ['Max Velocity', '101.7 mph', '97.7 mph', '105.5 mph'],
   ['Whiff% (Best Pitch)', '37.1% (ST)', '42.7% (SL)', '43.6% (CU)'],
]
for r, row_data in enumerate(data_rows, 1):
   for c, val in enumerate(row_data):
       table.rows[r].cells[c].text = val

doc.add_paragraph('')  # spacing

doc.add_paragraph(
   'Ohtani\'s best whiff pitch (sweeper, 37.1%) is outperformed by both Sanchez\'s '
   'slider (42.7%) and changeup (42.0%), as well as Misiorowski\'s curveball (43.6%). '
   'Overall whiff rates tell the same story: Ohtani (28.4%) trails Sanchez (29.7%) '
   'and Misiorowski (34.3%).'
)

# ---- Section 3: Plate Discipline ----
doc.add_heading('3. Plate Discipline and Control', level=2)
doc.add_paragraph(
   'The most significant gap between Ohtani and Sanchez is in command. Ohtani\'s walk '
   'rate of 7.6% is the highest among the three and nearly 60% higher than Sanchez\'s '
   'elite 4.8%. This directly impacts run prevention, as free baserunners increase '
   'scoring opportunities.'
)

table2 = doc.add_table(rows=4, cols=4)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Metric', 'Ohtani', 'Sanchez', 'Misiorowski']):
   table2.rows[0].cells[i].text = h
data2 = [
   ['K%', '27.9%', '27.6%', '39.6%'],
   ['BB%', '7.6%', '4.8%', '6.4%'],
   ['K-BB%', '20.3%', '22.8%', '33.2%'],
]
for r, row_data in enumerate(data2, 1):
   for c, val in enumerate(row_data):
       table2.rows[r].cells[c].text = val

doc.add_paragraph('')

# ---- Section 4: Run Expectancy ----
doc.add_heading('4. Run Expectancy Impact', level=2)
doc.add_paragraph(
   'The run expectancy delta (delta_pitcher_run_exp) directly measures how much each '
   'pitcher changes the expected runs on each pitch. A lower (more negative) value is better.'
)
doc.add_paragraph(
   'Sanchez (+14.487 total, +0.0080 per pitch) has been significantly better at preventing '
   'runs than Ohtani (+21.028 total, +0.0158 per pitch). This is despite Sanchez throwing '
   '465 more pitches. On a per-pitch basis, Sanchez gives up less than half the run value '
   'of Ohtani. This is a decisive finding.'
)

# ---- Section 5: Platoon ----
doc.add_heading('5. Platoon Splits', level=2)
doc.add_paragraph(
   'Ohtani struggles against left-handed batters (.545 OPS allowed) compared to Sanchez '
   '(.336) and Misiorowski (.288). Sanchez\'s 36.2% K% and 1.7% BB% against LHB are '
   'truly elite. While Ohtani handles RHB well (.383 OPS), his overall platoon profile '
   'is less dominant than his peers.'
)

# ---- Section 6: Conclusion ----
doc.add_heading('6. Conclusion', level=2)
doc.add_paragraph(
   'The data does not support the claim that Shohei Ohtani has been the best pitcher '
   'among these three NL aces in 2026.'
)
doc.add_paragraph(
   'Cristopher Sanchez has demonstrated superior command (4.8% BB%), better run prevention '
   'per pitch (+0.0080 RE delta), and a more effective pitch mix despite using only 3 pitch types. '
   'His sinker-changeup combination generates extreme ground ball tendencies (192 ground balls) '
   'and elite whiff rates on his secondary offerings (42.7% on slider, 42.0% on changeup).'
)
doc.add_paragraph(
   'Jacob Misiorowski has been the most dominant in terms of raw stuff, with a 39.6% K%, '
   '34.3% whiff rate, and 54.6% strike rate that place him in a separate tier. His 100.5 mph '
   'fastball with elite secondary pitches makes him arguably the hardest pitcher to face in the NL.'
)
doc.add_paragraph(
   'Ohtani remains an excellent pitcher — his contact suppression (zero barrels, 86.9 mph avg '
   'exit velocity) and expected stats are genuinely strong. However, his elevated walk rate, '
   'lower whiff rates, and significantly worse run expectancy impact suggest that his 2026 '
   'pitching performance has been outperformed by Sanchez and Misiorowski in the areas that '
   'matter most for run prevention.'
)

# Save
output_dir = os.path.join(os.path.dirname(__file__), '..', 'report')
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, 'MLB_Pitchers_Analysis_2026.docx')
doc.save(output_path)
print(f'Report saved to {output_path}')
